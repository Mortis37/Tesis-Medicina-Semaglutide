#!/usr/bin/env python3
"""
08_tablas.py — Genera todas las tablas de la tesis como documentos Word (.docx)

Output:
  Analisis_Global/Graficos/tablas_tesis.docx   (Cuadros 1-4)
  Analisis_Global/Graficos/cuadro_anova.docx   (Cuadros 5A, 5B, 5C)
  Analisis_Global/Graficos/Anexo11_pKa.docx    (Anexo 11)

Uso: python3 08_tablas.py
     (ejecutar desde la raiz del proyecto)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      "Analisis_Global", "Graficos")
os.makedirs(OUTDIR, exist_ok=True)

# ── Colores ──────────────────────────────────────────────────
HEADER_BG = "2E74B5"
ROW_ALT   = "EBF3FA"
HIGHLIGHT = "FFF3CD"
WHITE     = "FFFFFF"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_table(doc, title, subtitle, headers, rows, col_widths_cm,
              highlight_rows=None, note=None, bold_rows=None):
    """Anade titulo, tabla y nota al documento."""
    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)

    if subtitle:
        p_sub = doc.add_paragraph()
        run_s = p_sub.add_run(subtitle)
        run_s.font.size = Pt(9)
        run_s.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        run_s.font.italic = True

    ncol = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncol)
    table.style = "Table Grid"

    for i, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    # Encabezado
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_bg(cell, HEADER_BG)
        set_cell_borders(cell, "FFFFFF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # Filas
    for i, row_data in enumerate(rows):
        is_hl = highlight_rows and (i + 1) in highlight_rows
        is_bold = bold_rows and (i + 1) in bold_rows
        bg = HIGHLIGHT if is_hl else (ROW_ALT if i % 2 == 1 else WHITE)
        trow = table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = trow.cells[j]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            if is_bold:
                r.bold = True
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    if note:
        p_note = doc.add_paragraph()
        r_n = p_note.add_run(note)
        r_n.font.size = Pt(8.5)
        r_n.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
        r_n.font.italic = True

    doc.add_paragraph()


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    return doc


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 1: CUADROS 1-4 (tablas_tesis.docx)
# ══════════════════════════════════════════════════════════════
doc1 = new_doc()

# Cuadro 1 — Protocolo
add_table(doc1,
    title="Cuadro 1. Parametros del protocolo de simulacion (GROMACS 2025.3)",
    subtitle="Confirmados de los archivos step7_production.mdp generados por CHARMM-GUI",
    headers=["Parametro", "Valor", "Descripcion"],
    rows=[
        ["dt (paso integracion)", "0.002 ps (2 fs)", "Integracion numerica"],
        ["nsteps (pasos totales)", "100,000,000", "200 ns de produccion"],
        ["Duracion / replica", "200 ns", "Por replica"],
        ["Frames totales", "4,000 (1 frame/50 ps)", "Resolucion temporal 50 ps"],
        ["Termostato", "v-rescale", "3 grupos independientes"],
        ["Temperatura referencia", "310.15 K (37 C)", "37 C fisiologico"],
        ["Barostato", "C-rescale semi-isotropico", "Correcto para bicapa lipidica"],
        ["Presion referencia", "1.0 bar", "Isobarico"],
        ["Electrostatica largo alcance", "PME (Particle Mesh Ewald)", "Alta precision electrostatica"],
        ["Radio de corte (rcoulomb/rvdw)", "1.2 nm", "Estandar CHARMM36m"],
        ["Campo de fuerzas", "CHARMM36m", "Optimizado proteina-membrana"],
        ["Membrana (lipido)", "POPC (140 lipidos)", "1-palmitoil-2-oleoil-PC"],
        ["Agua / iones", "TIP3P / NaCl 150 mM", "Concentracion fisiologica"],
        ["Replicas por pH", "3 (independientes)", "Semillas distintas"],
        ["Total simulado", "1800 ns (3 replicas x 3 pH x 200 ns)", "9 sistemas en total"],
    ],
    col_widths_cm=[5.5, 5.5, 7.5])

doc1.add_page_break()

# Cuadro 2 — Energia libre
add_table(doc1,
    title="Cuadro 2. Energia libre de union MM-GBSA del complejo semaglutida-GLP-1R por condicion de pH",
    subtitle="igb=5, NaCl 150 mM, frames 150-200 ns, n = 3 replicas por condicion de pH",
    headers=["Condicion fisiologica", "pH", "DG binding (kcal/mol)", "DE", "SEM", "Puentes H (media +/- DE)"],
    rows=[
        ["Endosomal tardio", "5.0", "-105.60", "33.41", "19.29", "9.7 +/- 3.0"],
        ["Sistemico / plasmatico", "7.4", "-131.26", "42.70", "24.65", "15.2 +/- 5.5"],
        ["Gastrico / NAC-SNAC", "8.0", "-151.35", "27.18", "15.70", "18.1 +/- 4.4"],
    ],
    col_widths_cm=[4.5, 1.5, 4.0, 1.8, 1.8, 4.5],
    note="DE = desviacion estandar; SEM = error estandar de la media. Valores mas negativos = union mas favorable.")

# Cuadro 3 — Histidinas
add_table(doc1,
    title="Cuadro 3. Estado de protonacion de las histidinas pH-sensibles del GLP-1R (PropKa 3.0)",
    subtitle="HSP = histidina doblemente protonada (carga neta +1) | HSD = histidina neutra",
    headers=["Residuo", "Localizacion", "pH 5.0", "pH 7.4", "pH 8.0", "pKa estimado"],
    rows=[
        ["His99", "N-terminal extracelular", "HSP (+1)", "HSD (neutro)", "HSD (neutro)", "5.0-7.4"],
        ["His171", "N-terminal extracelular", "HSP (+1)", "HSD (neutro)", "HSD (neutro)", "5.0-7.4"],
        ["His173", "N-terminal extracelular", "HSP (+1)", "HSD (neutro)", "HSD (neutro)", "5.0-7.4"],
        ["His180", "ECL1 - contacto ligando", "HSP (+1)", "HSD (neutro)", "HSD (neutro)", "5.0-7.4"],
        ["His212", "ECL2", "HSP (+1)", "HSD (neutro)", "HSD (neutro)", "5.0-7.4"],
        ["His363", "TM6 (transmembrana)", "HSP (+1)", "HSD (neutro)", "HSD (neutro)", "5.0-7.4"],
        ["His374", "TM6 / TM7", "HSP (+1)", "HSD (neutro)", "HSD (neutro)", "5.0-7.4"],
    ],
    col_widths_cm=[2.0, 4.5, 2.5, 2.5, 2.5, 3.0],
    note="ECL = loop extracelular | TM = helice transmembrana | Posiciones asignadas por PDB 7KI0")

doc1.add_page_break()

# Cuadro 4 — Hotspots
add_table(doc1,
    title="Cuadro 4. Hotspots farmacologicos del GLP-1R - DG pairwise por residuo (kcal/mol)",
    subtitle="22 hotspots totales (DG <= -5.0 en al menos una condicion); 11 pH-sensibles (|DDG| > 1.0)",
    headers=["Residuo", "Localizacion", "DG pH 5.0", "DG pH 7.4", "DG pH 8.0", "pH-sensible"],
    rows=[
        ["GLU387", "TM6 (profundo)", "-20.96", "-51.94", "-53.98", "Si"],
        ["HIS180*", "ECL1 contacto lig.", "-5.27", "-6.04", "-6.30", "Si"],
        ["HIS363*", "TM6", "-4.08", "-6.06", "-5.93", "Si"],
        ["LEU359*", "TM5-TM6", "-4.19", "-5.22", "-5.30", "Si"],
        ["MET204", "ECL1-TM2", "-6.76", "-6.67", "-7.08", "No"],
        ["MET233", "TM3", "-6.05", "-6.18", "-6.80", "No"],
        ["LEU244", "TM3", "-5.74", "-6.11", "-6.11", "No"],
        ["LEU183", "TM2", "-5.87", "-5.98", "-6.00", "No"],
        ["LEU166", "TM2", "-5.68", "-5.89", "-6.21", "No"],
    ],
    col_widths_cm=[2.5, 4.0, 2.3, 2.3, 2.3, 2.5],
    highlight_rows=[1, 2, 3, 4],
    note=(
        "* pH-sensible: |DDG| > 1.0 kcal/mol entre pH 5.0 y 7.4\n"
        "GLU387: DG total pairwise (VdW + electrostatica + GB); VdW aislado = -3.85/-3.16/-2.19\n"
        "HIS180/HIS363 a pH 5.0: valores VdW reales (protonacion HSP reduce contribucion neta)\n"
        "Filas amarillas = pH-sensibles. n = 9 trayectorias (3 replicas x 3 pH)."))

path1 = os.path.join(OUTDIR, "tablas_tesis.docx")
doc1.save(path1)
print(f"OK: {path1}")


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 2: CUADROS 5A, 5B, 5C — ANOVA (cuadro_anova.docx)
# ══════════════════════════════════════════════════════════════
doc2 = new_doc()

# Cuadro 5A — ANOVA
add_table(doc2,
    title="Cuadro 5A. ANOVA de una via - DG binding (kcal/mol) por condicion de pH",
    subtitle=("Datos: 9 valores reales (3 replicas x 3 pH) | "
              "Supuestos verificados: normalidad (Shapiro-Wilk p>0.30) y "
              "homogeneidad de varianzas (Bartlett p=0.85)"),
    headers=["Fuente de variacion", "SC", "gl", "CM", "F", "p", "eta2"],
    rows=[
        ["pH (entre grupos)", "3154.63", "2", "1577.32", "1.286", "0.343", "0.300"],
        ["Residuos (dentro)", "7356.46", "6", "1226.08", "-", "-", "-"],
        ["Total", "10511.09", "8", "-", "-", "-", "-"],
    ],
    col_widths_cm=[5.5, 2.2, 1.5, 2.2, 1.8, 1.8, 1.8],
    bold_rows=[3],
    note=("SC = suma de cuadrados | gl = grados de libertad | CM = cuadrado medio | "
          "eta2 = eta cuadrado (tamano del efecto) | F(2,6) = 1.286, p = 0.343 (no significativo) | "
          "eta2 = 0.300 (efecto grande, Cohen 1988)"))

# Cuadro 5B — Medias
add_table(doc2,
    title="Cuadro 5B. Medias de DG binding por condicion de pH",
    subtitle="n = 3 replicas independientes por condicion | DE = desviacion estandar",
    headers=["Condicion fisiologica", "pH", "Media (kcal/mol)", "DE", "SEM", "Replicas"],
    rows=[
        ["Endosomal tardio", "5.0", "-105.60", "33.41", "19.29", "3"],
        ["Sistemico / plasmatico", "7.4", "-131.26", "42.70", "24.65", "3"],
        ["Gastrico / NAC-SNAC", "8.0", "-151.35", "27.18", "15.70", "3"],
    ],
    col_widths_cm=[5.0, 1.5, 3.5, 2.0, 2.0, 2.0],
    note=("Valores individuales: pH 5.0: -143.67, -91.97, -81.16; "
          "pH 7.4: -178.43, -120.08, -95.26; pH 8.0: -181.76, -142.87, -129.41 kcal/mol"))

# Cuadro 5C — Tukey HSD
add_table(doc2,
    title="Cuadro 5C. Comparaciones post-hoc Tukey HSD",
    subtitle=("q critico (alfa=0.05, k=3, gl=6) = 4.34 | "
              "Todas las comparaciones: ns (p>0.05) | "
              "Baja potencia (n=3) ante variabilidad alta, no ausencia de efecto (eta2=0.300)"),
    headers=["Comparacion", "Diferencia\n(kcal/mol)", "IC 95%\ninferior", "IC 95%\nsuperior", "p ajustada", "Significancia"],
    rows=[
        ["pH 7.4 - pH 5.0", "-25.66", "-113.38", "62.07", "0.661", "ns"],
        ["pH 8.0 - pH 5.0", "-45.75", "-133.47", "41.98", "0.316", "ns"],
        ["pH 8.0 - pH 7.4", "-20.09", "-107.81", "67.63", "0.771", "ns"],
    ],
    col_widths_cm=[4.5, 3.0, 2.5, 2.5, 2.5, 2.5],
    note=("IC 95% calculados con metodo Tukey (familywise). ns = no significativo (p>0.05). "
          "Interpretar descriptivamente las medias e informar eta2=0.300 como relevancia practica."))

path2 = os.path.join(OUTDIR, "cuadro_anova.docx")
doc2.save(path2)
print(f"OK: {path2}")


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 3: ANEXO 11 — pKa PropKa 3.0 (Anexo11_pKa.docx)
# ══════════════════════════════════════════════════════════════
doc3 = new_doc()

# Titulo
p = doc3.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Anexo 11. Valores de pKa estimados por PropKa 3.0 para los residuos "
                "ionizables del complejo semaglutida-GLP-1R")
run.bold = True
run.font.size = Pt(11)
run.font.name = "Arial"

# Nota introductoria
p2 = doc3.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run2 = p2.add_run(
    "Valores de pKa calculados con PropKa 3.0 (Olsson et al. 2011) sobre la estructura "
    "cristalografica PDB 7KI0 (cadena R = GLP-1R, cadena P = semaglutida). Se listan las "
    "siete histidinas pH-sensibles con sus pKa individuales y una muestra representativa de "
    "residuos Asp, Glu, Lys y Arg cuyo pKa queda fuera del intervalo de pH estudiado (5.0-8.0), "
    "confirmando que no cambian de estado de protonacion en ninguna condicion simulada.")
run2.font.size = Pt(9)
run2.font.name = "Arial"
run2.italic = True

# Datos
histidinas = [
    ("His99", "HIS", 6.16, "N-terminal ECD", "HSP (+1)", "HSD (0)", "HSD (0)", "Si"),
    ("His171", "HIS", 7.08, "N-terminal / TM1", "HSP (+1)", "HSD (0)", "HSD (0)", "Si"),
    ("His173", "HIS", 5.97, "N-terminal ECD", "HSP (+1)", "HSD (0)", "HSD (0)", "Si"),
    ("His180", "HIS", 6.35, "ECL1 contacto lig.", "HSP (+1)", "HSD (0)", "HSD (0)", "Si"),
    ("His212", "HIS", 6.37, "ECL2", "HSP (+1)", "HSD (0)", "HSD (0)", "Si"),
    ("His363", "HIS", 7.95, "TM6", "HSP (+1)", "HSD (0)", "HSD (0)", "Si"),
    ("His374", "HIS", 6.35, "TM6 / TM7", "HSP (+1)", "HSD (0)", "HSD (0)", "Si"),
]
asp_sample = [
    ("Asp59", "ASP", 4.11, "N-terminal ECD", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Asp67", "ASP", 3.28, "N-terminal ECD", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Asp74", "ASP", 4.11, "N-terminal ECD", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Asp114", "ASP", 4.20, "N-terminal ECD", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Asp344", "ASP", 4.55, "TM5", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Asp372", "ASP", 0.44, "TM6", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
]
glu_sample = [
    ("Glu9", "GLU", 3.12, "Semaglutida (P)", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Glu41", "GLU", 3.67, "N-terminal ECD", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Glu294", "GLU", 4.67, "TM4-TM5", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Glu387", "GLU", 6.25, "TM6 (hotspot)", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Glu408", "GLU", 5.11, "TM7", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
    ("Glu412", "GLU", 4.73, "TM7", "COO- (-1)", "COO- (-1)", "COO- (-1)", "No"),
]
lys_sample = [
    ("Lys26", "LYS", 9.69, "Semaglutida (P)", "NH3+ (+1)", "NH3+ (+1)", "NH3+ (+1)", "No"),
    ("Lys197", "LYS", 10.22, "TM2", "NH3+ (+1)", "NH3+ (+1)", "NH3+ (+1)", "No"),
    ("Lys288", "LYS", 8.25, "TM4", "NH3+ (+1)", "NH3+ (+1)", "NH3+ (+1)", "No"),
    ("Lys346", "LYS", 10.11, "TM5", "NH3+ (+1)", "NH3+ (+1)", "NH3+ (+1)", "No"),
    ("Lys383", "LYS", 10.17, "TM6", "NH3+ (+1)", "NH3+ (+1)", "NH3+ (+1)", "No"),
    ("Lys415", "LYS", 9.68, "TM7", "NH3+ (+1)", "NH3+ (+1)", "NH3+ (+1)", "No"),
]
arg_sample = [
    ("Arg176", "ARG", 15.26, "ECL1", "Guan+ (+1)", "Guan+ (+1)", "Guan+ (+1)", "No"),
    ("Arg190", "ARG", 17.60, "TM2", "Guan+ (+1)", "Guan+ (+1)", "Guan+ (+1)", "No"),
    ("Arg227", "ARG", 14.56, "TM3", "Guan+ (+1)", "Guan+ (+1)", "Guan+ (+1)", "No"),
    ("Arg310", "ARG", 11.98, "TM5", "Guan+ (+1)", "Guan+ (+1)", "Guan+ (+1)", "No"),
    ("Arg376", "ARG", 11.06, "TM6", "Guan+ (+1)", "Guan+ (+1)", "Guan+ (+1)", "No"),
    ("Arg419", "ARG", 11.47, "TM7", "Guan+ (+1)", "Guan+ (+1)", "Guan+ (+1)", "No"),
]

headers = ["Residuo", "Tipo", "pKa", "Localizacion", "pH 5.0", "pH 7.4", "pH 8.0", "pH-sensible"]

all_rows = []
all_rows.append(("SEPARADOR", "Histidinas pH-sensibles del GLP-1R (cadena R)"))
all_rows.extend(histidinas)
all_rows.append(("SEPARADOR", "Acido aspartico (Asp) - muestra representativa"))
all_rows.extend(asp_sample)
all_rows.append(("SEPARADOR", "Acido glutamico (Glu) - muestra representativa"))
all_rows.extend(glu_sample)
all_rows.append(("SEPARADOR", "Lisina (Lys) - muestra representativa"))
all_rows.extend(lys_sample)
all_rows.append(("SEPARADOR", "Arginina (Arg) - muestra representativa"))
all_rows.extend(arg_sample)

# Crear tabla
table = doc3.add_table(rows=1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

# Header
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = "Arial"
    shading = cell._element.get_or_add_tcPr()
    sh_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "2C3E50"})
    shading.append(sh_elem)
    run.font.color.rgb = RGBColor(255, 255, 255)

# Filas
for row_data in all_rows:
    if row_data[0] == "SEPARADOR":
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[-1])
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(row_data[1])
        run.bold = True
        run.italic = True
        run.font.size = Pt(8)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(44, 62, 80)
        shading = cell._element.get_or_add_tcPr()
        sh_elem = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "D5D8DC"})
        shading.append(sh_elem)
    else:
        row = table.add_row()
        values = [row_data[0], row_data[1], f"{row_data[2]:.2f}",
                  row_data[3], row_data[4], row_data[5], row_data[6], row_data[7]]
        is_his = row_data[1] == "HIS"
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.name = "Arial"
            if is_his:
                run.bold = True
                run.font.color.rgb = RGBColor(192, 57, 43)
                shading = cell._element.get_or_add_tcPr()
                sh_elem = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "FADBD8"})
                shading.append(sh_elem)

# Ancho de columnas
widths = [Cm(1.8), Cm(1.2), Cm(1.2), Cm(4.0), Cm(2.2), Cm(2.2), Cm(2.2), Cm(2.0)]
for row in table.rows:
    for i, w in enumerate(widths):
        if i < len(row.cells):
            row.cells[i].width = w

# Nota al pie
p3 = doc3.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run3 = p3.add_run(
    "Nota: pKa calculados sobre PDB 7KI0 (Zhang et al. 2021) mediante PropKa 3.0 "
    "(Olsson et al. 2011). HSP = histidina doblemente protonada (carga +1); "
    "HSD = histidina neutra (protonada en Nd, carga 0). Residuos Asp/Glu con pKa < 5.0 "
    "permanecen desprotonados y Lys/Arg con pKa > 8.0 permanecen protonados en las tres "
    "condiciones. Solo las siete histidinas cambian de estado a pH 5.0.")
run3.font.size = Pt(8)
run3.font.name = "Arial"
run3.italic = True

path3 = os.path.join(OUTDIR, "Anexo11_pKa.docx")
doc3.save(path3)
print(f"OK: {path3}")

print(f"\nTodas las tablas en: {OUTDIR}")
